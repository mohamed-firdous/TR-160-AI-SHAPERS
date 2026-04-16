from __future__ import annotations

import io
import re
from dataclasses import dataclass

import fitz
from docx import Document


@dataclass
class ParsedDocument:
    filename: str
    raw_text: str
    paragraphs: list[str]


def parse_uploaded_file(filename: str, file_bytes: bytes) -> ParsedDocument:
    extension = filename.lower().rsplit(".", 1)[-1]
    stream = io.BytesIO(file_bytes)

    if extension == "pdf":
        raw_text = _read_pdf(stream)
    elif extension == "docx":
        raw_text = _read_docx(stream)
    elif extension == "txt":
        raw_text = _read_txt(stream)
    else:
        raise ValueError("Unsupported file type. Use PDF, DOCX, or TXT.")

    clean = _clean_text(raw_text)
    paragraphs = _split_paragraphs(clean)

    if not paragraphs:
        raise ValueError("No meaningful paragraphs found in file.")

    return ParsedDocument(filename=filename, raw_text=clean, paragraphs=paragraphs)


def _read_pdf(stream: io.BytesIO) -> str:
    document = fitz.open(stream=stream.read(), filetype="pdf")
    return "\n".join(page.get_text("text") for page in document)


def _read_docx(stream: io.BytesIO) -> str:
    document = Document(stream)
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def _read_txt(stream: io.BytesIO) -> str:
    payload = stream.read()
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode TXT file.")


def _clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", text)
    return text.strip()


def _split_paragraphs(text: str, min_chars: int = 60) -> list[str]:
    blocks = [item.strip() for item in re.split(r"\n\s*\n", text) if item.strip()]
    paragraphs: list[str] = []
    buffer = ""

    for block in blocks:
        if len(block) < min_chars:
            buffer = f"{buffer} {block}".strip()
            continue
        if buffer:
            block = f"{buffer} {block}".strip()
            buffer = ""
        paragraphs.append(block)

    if buffer:
        if paragraphs:
            paragraphs[-1] = f"{paragraphs[-1]} {buffer}".strip()
        else:
            paragraphs.append(buffer)

    return paragraphs
